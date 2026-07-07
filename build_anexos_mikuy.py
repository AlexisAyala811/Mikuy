from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DOC = Path(r"D:\Documentación de Paginas\Mikuy\Documento de Mikuy (1).docx")
OUT_DOC = Path(r"D:\Paginas de Internet\Mikuy\Documento_Mikuy_con_anexos.docx")

APA_TABLE_COUNTER = None
APA_TABLE_TITLES = []


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge_name, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        if value is None:
            edge.set(qn("w:val"), "nil")
        else:
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), str(value.get("size", 8)))
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), value.get("color", "000000"))


def clear_cell_shading(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is not None:
        tc_pr.remove(shd)


def format_run(run, bold=None, italic=None, size=12, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text="", style=None, align=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        format_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        format_run(r2)
    else:
        r = p.add_run(text)
        format_run(r)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    format_run(r, bold=True, size=13 if level <= 2 else 12)
    if level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def format_table(table, widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for style_name in ("Table Grid", "Cuadrícula de tabla", "Tabla con cuadrícula"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cell, widths[col_idx])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    format_run(run, size=9 if len(cell.text) > 140 else 10)
            if row_idx == 0:
                set_cell_shading(cell, "E8D9C8")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        format_run(run, bold=True, size=10)
    set_repeat_table_header(table.rows[0])


def add_table(doc, headers, rows, widths=None):
    if APA_TABLE_COUNTER is not None and APA_TABLE_TITLES:
        return add_apa_table(doc, APA_TABLE_COUNTER, APA_TABLE_TITLES.pop(0), headers, rows, widths)

    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value
    format_table(table, widths)
    doc.add_paragraph()
    return table


def add_apa_table(doc, counter, title, headers, rows, widths=None):
    number = counter["value"]
    counter["value"] += 1

    p_num = doc.add_paragraph()
    r_num = p_num.add_run(f"Tabla A.{number}")
    format_run(r_num, bold=True)
    p_num.paragraph_format.space_after = Pt(0)

    p_title = doc.add_paragraph()
    r_title = p_title.add_run(title)
    format_run(r_title, italic=True)
    p_title.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            clear_cell_shading(cell)
            if widths:
                set_cell_width(cell, widths[col_idx])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    format_run(run, size=9 if len(cell.text) > 140 else 10, bold=(row_idx == 0))

            set_cell_borders(cell, left=None, right=None, top=None, bottom=None)
            if row_idx == 0:
                set_cell_borders(
                    cell,
                    top={"size": 12, "color": "000000"},
                    bottom={"size": 8, "color": "000000"},
                    left=None,
                    right=None,
                )
            if row_idx == len(table.rows) - 1 and row_idx != 0:
                set_cell_borders(
                    cell,
                    top=None,
                    bottom={"size": 12, "color": "000000"},
                    left=None,
                    right=None,
                )

    set_repeat_table_header(table.rows[0])
    doc.add_paragraph()
    return table


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("Nota. ")
    format_run(r, bold=True, italic=True)
    r2 = p.add_run(text)
    format_run(r2, italic=True)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_user_story_detail(doc, counter, number, story):
    rows = [
        ("ID", story["id"]),
        ("Nombre", story["nombre"]),
        ("Módulo / Sprint", story["modulo"]),
        ("Prioridad", story["prioridad"]),
        ("Descripción", story["descripcion"]),
        ("Criterios de Aceptación", story["criterios"]),
        ("Reglas de Negocio", story["reglas"]),
        ("Estimación", story["estimacion"]),
        ("Desglose de Tareas Técnicas", story["tareas"]),
        ("Evidencia en el Proyecto", story["evidencia"]),
    ]
    add_apa_table(
        doc,
        counter,
        f"Historia de Usuario {number:02d}: {story['nombre']}",
        ["Campo", "Detalle"],
        rows,
        [4.0, 12.8],
    )
    add_note(doc, "La historia de usuario fue definida a partir de los requerimientos funcionales del sistema de reservas Mikuy. Elaboración propia.")


def switch_section(doc, orientation):
    section = doc.add_section()
    section.orientation = orientation
    if orientation == WD_ORIENT.LANDSCAPE:
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    else:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return section


def build():
    global APA_TABLE_COUNTER, APA_TABLE_TITLES

    doc = Document(BASE_DOC)

    doc.add_page_break()
    add_heading(doc, "Anexo", 2)
    add_paragraph(
        doc,
        "La presente sección reúne los instrumentos, matrices y evidencias técnicas que sustentan el desarrollo de la plataforma web de reservas y gestión de mesas para el restaurante Mikuy. Los anexos se organizan conforme al modelo del documento de referencia, priorizando la trazabilidad entre problema, objetivos, variables, sprints, validación funcional y evaluación de usabilidad.",
    )

    annex_table_counter = {"value": 1}
    APA_TABLE_COUNTER = annex_table_counter
    APA_TABLE_TITLES = [
        "Ficha de An\u00e1lisis Documental para Validaci\u00f3n Funcional",
        "Cuestionario System Usability Scale (SUS)",
        "Resultados Consolidados del Cuestionario SUS",
        "Resumen del Product Backlog del Proyecto Mikuy",
        "Entregables de Ingenier\u00eda de Software del Proyecto Mikuy",
        "Evidencia de Control de Versiones",
    ]

    switch_section(doc, WD_ORIENT.LANDSCAPE)
    add_heading(doc, "Anexo 1. Matriz de Consistencia", 3)
    add_paragraph(
        doc,
        "Título: Plataforma web para la gestión automatizada de reservas y mesas del restaurante Mikuy, Huamanga, 2026."
    )
    matrix_rows = [
        (
            "Problema general\n¿Cuáles son los resultados del desarrollo de una plataforma web para optimizar la gestión de reservas y mesas del restaurante Mikuy?\n\nProblemas específicos\n1. ¿Cuáles son los requerimientos funcionales del proceso de reserva?\n2. ¿Qué arquitectura y diseño permiten organizar la solución?\n3. ¿Cuáles son los resultados de la implementación por sprints?\n4. ¿Cómo se valida el funcionamiento de los módulos principales?\n5. ¿Cuál es el nivel de usabilidad percibido por los usuarios piloto?",
            "Objetivo general\nDesarrollar una plataforma web para automatizar la gestión de reservas y mesas del restaurante Mikuy.\n\nObjetivos específicos\n1. Determinar los requerimientos del proceso de reserva.\n2. Diseñar la arquitectura, base de datos e interfaces principales.\n3. Implementar los módulos mediante sprints.\n4. Validar el funcionamiento técnico del sistema.\n5. Evaluar la usabilidad mediante el cuestionario SUS.",
            "Variable de interés\nX: Plataforma web de reservas y mesas.\n\nDimensiones\nX1: Análisis de requerimientos.\nX2: Diseño arquitectónico y de interfaz.\nX3: Implementación del sistema.\nX4: Funcionamiento del software.\nX5: Usabilidad de la plataforma.",
            "Tipo: Aplicada.\nNivel: Descriptivo.\nDiseño: No experimental.\nPoblación: clientes potenciales y personal administrativo del restaurante Mikuy.\nMuestra: 30 usuarios piloto y 3 trabajadores.\nTécnicas: entrevista, análisis documental y encuesta.\nInstrumentos: guía de entrevista, ficha de análisis documental y cuestionario SUS.",
        )
    ]
    add_apa_table(
        doc,
        annex_table_counter,
        "Matriz de Consistencia del Proyecto Mikuy",
        ["Problemas", "Objetivos", "Variables", "Método de investigación"],
        matrix_rows,
        [7.2, 7.2, 5.2, 6.2],
    )
    add_note(doc, "La matriz articula los componentes metodológicos que orientan el desarrollo y validación de la plataforma Mikuy. Elaboración propia.")

    switch_section(doc, WD_ORIENT.PORTRAIT)
    add_heading(doc, "Anexo 2. Instrumentos de Recolección de Datos", 3)
    add_heading(doc, "2.1. Guía de entrevista al personal del restaurante", 4)
    interview_questions = [
        "¿Cómo se registran actualmente las reservas y qué problemas se presentan durante horas de alta demanda?",
        "¿Qué datos mínimos debe registrar el sistema para considerar válida una solicitud de reserva?",
        "¿Qué reglas de negocio se aplican para asignar mesas según la cantidad de personas?",
        "¿Qué acciones debe poder realizar el administrador sobre una reserva pendiente, confirmada o cancelada?",
        "¿Qué indicadores operativos necesita visualizar el personal para controlar aforo, agenda diaria y disponibilidad?",
        "¿Qué canales de comunicación deben usarse para confirmar o consultar el estado de una reserva?",
    ]
    for question in interview_questions:
        add_paragraph(doc, question, style="List Number")

    add_heading(doc, "2.2. Ficha de análisis documental para validación funcional", 4)
    checklist_rows = [
        ("Gestión de acceso", "Login administrativo", "Verificar autenticación y protección de rutas internas mediante cookies.", "Cumple / No cumple"),
        ("Seguridad", "Hashing PBKDF2", "Comprobar que las contraseñas se almacenen con sal, SHA-256 e iteraciones configuradas.", "Cumple / No cumple"),
        ("Reservas", "Wizard público", "Validar registro de fecha, hora, cantidad de personas y datos de contacto.", "Cumple / No cumple"),
        ("Aforo", "Asignación Best-Fit", "Comprobar que el sistema elija la mesa activa de menor capacidad suficiente.", "Cumple / No cumple"),
        ("Comunicación", "WhatsApp", "Verificar generación de enlace para confirmación o consulta de reserva.", "Cumple / No cumple"),
        ("Administración", "Dashboard", "Validar indicadores de pendientes, ocupación, agenda y acciones rápidas.", "Cumple / No cumple"),
    ]
    add_table(doc, ["Dimensión", "Ítem", "Criterio de verificación", "Resultado"], checklist_rows, [3.3, 3.2, 8.1, 3.4])

    add_heading(doc, "2.3. Cuestionario System Usability Scale (SUS)", 4)
    sus_rows = [
        ("1", "Creo que me gustaría usar esta plataforma web frecuentemente."),
        ("2", "Encontré la plataforma innecesariamente compleja."),
        ("3", "Consideré que la plataforma fue fácil de usar."),
        ("4", "Creo que necesitaría ayuda técnica para utilizar la plataforma."),
        ("5", "Encontré que las funciones de reserva estuvieron bien integradas."),
        ("6", "Consideré que hubo demasiada inconsistencia en la plataforma."),
        ("7", "Imagino que la mayoría de usuarios aprendería a usar la plataforma rápidamente."),
        ("8", "Encontré la plataforma difícil de utilizar."),
        ("9", "Me sentí seguro al usar el sistema de reservas."),
        ("10", "Necesité aprender muchas cosas antes de poder usar la plataforma."),
    ]
    add_table(doc, ["Ítem", "Enunciado"], sus_rows, [1.6, 14.8])
    add_note(doc, "Los ítems impares se interpretan en sentido positivo y los pares en sentido inverso, siguiendo la metodología SUS de Brooke.")

    add_page_break = doc.add_page_break
    add_page_break()
    add_heading(doc, "Anexo 3. Resultados Estadísticos de Usabilidad", 3)
    add_paragraph(
        doc,
        "Los resultados consolidados del cuestionario SUS aplicado a 30 usuarios piloto permitieron obtener un puntaje global de 88.50, clasificado como un nivel de usabilidad aceptable y satisfactorio. La siguiente matriz resume los valores empleados para el cálculo presentado en el Capítulo IV."
    )
    sus_result_rows = [
        ("1", "4.30", "3.30", "Contribución positiva"),
        ("2", "1.70", "3.30", "Contribución inversa"),
        ("3", "4.40", "3.40", "Contribución positiva"),
        ("4", "1.60", "3.40", "Contribución inversa"),
        ("5", "4.50", "3.50", "Contribución positiva"),
        ("6", "1.50", "3.50", "Contribución inversa"),
        ("7", "4.60", "3.60", "Contribución positiva"),
        ("8", "1.50", "3.50", "Contribución inversa"),
        ("9", "4.70", "3.70", "Contribución positiva"),
        ("10", "1.70", "3.30", "Contribución inversa"),
    ]
    add_table(doc, ["Ítem", "Promedio Likert", "Contribución SUS", "Interpretación"], sus_result_rows, [1.5, 3.6, 3.6, 7.6])
    add_note(doc, "La suma de contribuciones fue de 35.40 puntos; al multiplicarse por 2.5 se obtuvo un puntaje SUS de 88.50. Elaboración propia.")

    add_page_break()
    add_heading(doc, "Anexo 4. Detalle de Historias de Usuario del Product Backlog", 3)
    backlog_rows = [
        ("HU-01", "Registro de clientes", "Como comensal, quiero registrar mis datos para solicitar una reserva.", "Alta", "5"),
        ("HU-02", "Autenticación administrativa", "Como administrador, quiero iniciar sesión de forma segura para gestionar la operación.", "Alta", "5"),
        ("HU-03", "Gestión de mesas", "Como administrador, quiero crear y actualizar mesas para controlar el aforo disponible.", "Alta", "5"),
        ("HU-04", "Gestión de platos", "Como administrador, quiero mantener el catálogo gastronómico visible para los clientes.", "Media", "3"),
        ("HU-05", "Wizard público de reserva", "Como cliente, quiero reservar en pasos claros indicando fecha, hora y personas.", "Alta", "8"),
        ("HU-06", "Asignación Best-Fit", "Como sistema, quiero asignar la mesa activa de menor capacidad suficiente.", "Alta", "8"),
        ("HU-07", "Confirmación y consulta", "Como cliente, quiero consultar y descargar el comprobante de mi reserva.", "Media", "5"),
        ("HU-08", "Dashboard administrativo", "Como administrador, quiero visualizar pendientes, agenda y ocupación del día.", "Alta", "5"),
        ("HU-09", "Cancelación y liberación de mesa", "Como usuario autorizado, quiero cancelar reservas sin bloquear la mesa.", "Alta", "5"),
        ("HU-10", "Notificación por WhatsApp", "Como personal operativo, quiero contactar al cliente mediante enlace generado.", "Media", "3"),
    ]
    add_table(doc, ["ID", "Historia de usuario", "Descripción", "Prioridad", "Puntos"], backlog_rows, [1.7, 4.0, 8.2, 2.3, 1.8])
    add_note(doc, "El backlog sintetiza los requerimientos implementados durante los sprints del proyecto Mikuy. Elaboración propia.")

    detailed_stories = [
        {
            "id": "HU-01",
            "nombre": "Registro de Clientes y Validación de Datos",
            "modulo": "Sprint 1 - Gestión de clientes y base de datos",
            "prioridad": "Alta",
            "descripcion": "Como cliente del restaurante Mikuy,\nQuiero registrar mis datos personales básicos en la plataforma,\nPara poder solicitar reservas y recibir confirmaciones asociadas a mi nombre, correo y teléfono.",
            "criterios": "1. El formulario debe solicitar nombre completo, correo electrónico y teléfono.\n2. Los campos obligatorios deben validarse antes de persistir el registro.\n3. El correo debe almacenarse en minúsculas y sin espacios laterales para evitar duplicidad lógica.\n4. El teléfono debe normalizarse conservando únicamente dígitos para facilitar consultas posteriores.\n5. Si el correo ya existe, el sistema debe actualizar los datos básicos del cliente sin crear un duplicado.\n6. El cliente registrado debe quedar asociado automáticamente a la reserva creada desde el flujo público.",
            "reglas": "RN-01: Un cliente se identifica principalmente por su correo electrónico.\nRN-02: El teléfono se usa como dato de contacto y como criterio alternativo de consulta.\nRN-03: No se debe crear una reserva si los datos mínimos del cliente son inválidos.",
            "estimacion": "5 puntos de historia / Aprox. 8 horas ideales.",
            "tareas": "T1 (Dominio): Definir la entidad Cliente con nombre, teléfono, correo y relación con reservas.\nT2 (Persistencia): Configurar el mapeo relacional en Entity Framework Core y restricciones de longitud.\nT3 (Backend): Implementar validaciones en DTOs, ViewModels y controladores.\nT4 (Frontend): Diseñar vistas de registro, edición, detalle y listado de clientes.\nT5 (Integración): Reutilizar el cliente existente durante la creación de reserva pública.\nT6 (Pruebas): Validar creación, edición, normalización de correo/teléfono y asociación con reserva.",
            "evidencia": "Reserva.Domain/Entities/Cliente.cs; Reserva.Web/Controllers/ClientesController.cs; Reserva.Web/Views/Clientes; Reserva.Web/Controllers/ReservasController.cs.",
        },
        {
            "id": "HU-02",
            "nombre": "Autenticación Administrativa Segura",
            "modulo": "Sprint 1 - Seguridad administrativa",
            "prioridad": "Alta",
            "descripcion": "Como administrador de Mikuy,\nQuiero iniciar sesión mediante credenciales protegidas,\nPara acceder únicamente a las funciones internas de gestión del restaurante.",
            "criterios": "1. El sistema debe mostrar un formulario de inicio de sesión para usuarios administrativos.\n2. Las contraseñas deben almacenarse con formato PBKDF2, iteraciones, sal y clave derivada en Base64.\n3. La verificación de contraseña debe recalcular PBKDF2 usando la sal almacenada.\n4. La comparación de claves debe ejecutarse con FixedTimeEquals para evitar filtraciones por tiempo.\n5. Las rutas administrativas deben requerir autenticación mediante cookies.\n6. El sistema debe rechazar credenciales vacías, incorrectas o hashes mal formados sin lanzar excepciones visibles al usuario.",
            "reglas": "RN-01: Ninguna contraseña administrativa puede almacenarse en texto plano.\nRN-02: Solo usuarios autenticados pueden acceder al dashboard, CRUDs y acciones de cambio de estado.\nRN-03: El cierre de sesión debe invalidar la cookie de autenticación.",
            "estimacion": "5 puntos de historia / Aprox. 8 horas ideales.",
            "tareas": "T1 (Seguridad): Implementar PasswordHashing.Hash y PasswordHashing.Verify con PBKDF2-SHA256.\nT2 (Persistencia): Hashear usuarios semilla y actualizar contraseñas antiguas en DataSeeder.\nT3 (Backend): Programar CuentaController para login, creación de claims y logout.\nT4 (Configuración): Registrar autenticación por cookies y autorización en Program.cs.\nT5 (Interfaz): Crear vista de login con validación visual de errores.\nT6 (Pruebas): Verificar login correcto, rechazo de credenciales erróneas y protección de rutas.",
            "evidencia": "Reserva.Infrastructure/Security/PasswordHashing.cs; Reserva.Infrastructure/Persistence/DataSeeder.cs; Reserva.Web/Controllers/CuentaController.cs; Reserva.Web/Program.cs.",
        },
        {
            "id": "HU-03",
            "nombre": "Gestión Administrativa de Mesas",
            "modulo": "Sprint 2 - Gestión de aforo",
            "prioridad": "Alta",
            "descripcion": "Como administrador del restaurante,\nQuiero registrar y actualizar las mesas disponibles con su capacidad y ubicación,\nPara controlar el aforo operativo y permitir asignaciones automáticas de reservas.",
            "criterios": "1. El sistema debe listar mesas mostrando número, capacidad, ubicación y estado.\n2. El administrador debe poder crear nuevas mesas activas con capacidad mayor que cero.\n3. El administrador debe poder editar capacidad, ubicación y estado operativo de una mesa.\n4. Las mesas inactivas no deben aparecer como opción disponible en nuevas reservas.\n5. Las mesas con reservas vigentes en la misma fecha y hora deben excluirse de la disponibilidad.\n6. La cantidad total de mesas configuradas debe servir como referencia para el control de aforo.",
            "reglas": "RN-01: Una mesa inactiva representa una mesa fuera de servicio y no participa en asignaciones.\nRN-02: La capacidad define el número máximo de personas que pueden ser asignadas a la mesa.\nRN-03: El número de mesa debe ser comprensible para el personal operativo.",
            "estimacion": "5 puntos de historia / Aprox. 8 horas ideales.",
            "tareas": "T1 (Dominio): Definir Mesa con número, capacidad, ubicación, estado e identificador.\nT2 (Persistencia): Configurar entidad y datos semilla de mesas en SQL Server.\nT3 (Backend): Desarrollar acciones Index, Create, Edit y validaciones en MesasController.\nT4 (Frontend): Crear vistas administrativas y formulario parcial _MesaForm.\nT5 (Integración): Conectar mesas activas con formularios de reserva y algoritmo Best-Fit.\nT6 (Pruebas): Validar creación, edición, desactivación y consulta de disponibilidad.",
            "evidencia": "Reserva.Domain/Entities/Mesa.cs; Reserva.Web/Controllers/MesasController.cs; Reserva.Web/Views/Mesas; Reserva.Infrastructure/Persistence/DataSeeder.cs.",
        },
        {
            "id": "HU-04",
            "nombre": "Gestión de Platos del Restaurante",
            "modulo": "Sprint 2 - Catálogo gastronómico",
            "prioridad": "Media",
            "descripcion": "Como administrador,\nQuiero administrar los platos representativos del restaurante,\nPara mantener actualizada la información gastronómica visible para los clientes.",
            "criterios": "1. El sistema debe listar platos registrados con nombre, categoría, precio y estado.\n2. Cada plato debe registrar descripción e imagen representativa.\n3. El administrador debe poder crear platos nuevos y editar platos existentes.\n4. Solo los platos activos deben formar parte de la visualización pública.\n5. La carga de imágenes debe limitarse a extensiones permitidas y rutas controladas.\n6. El listado público debe mantener coherencia visual con la identidad gastronómica de Mikuy.",
            "reglas": "RN-01: Un plato inactivo no debe promocionarse en la página pública.\nRN-02: El precio debe ser un valor numérico válido.\nRN-03: Las imágenes deben almacenarse en una ruta pública controlada del proyecto.",
            "estimacion": "3 puntos de historia / Aprox. 5 horas ideales.",
            "tareas": "T1 (Dominio): Configurar entidad Plato con propiedades gastronómicas.\nT2 (Backend): Implementar acciones administrativas y validaciones en PlatosController.\nT3 (Frontend): Diseñar formularios Create/Edit, listado y vista pública del menú.\nT4 (Archivos): Validar extensiones permitidas y almacenamiento en wwwroot/img/platos.\nT5 (Datos): Sembrar platos representativos de la gastronomía ayacuchana.\nT6 (Pruebas): Verificar carga, edición, activación y visualización pública.",
            "evidencia": "Reserva.Domain/Entities/Plato.cs; Reserva.Web/Controllers/PlatosController.cs; Reserva.Web/Views/Platos; Reserva.Web/wwwroot/img/platos.",
        },
        {
            "id": "HU-05",
            "nombre": "Formulario Wizard Público de Reserva",
            "modulo": "Sprint 3 - Reserva pública",
            "prioridad": "Alta",
            "descripcion": "Como comensal,\nQuiero completar mi reserva mediante un asistente dividido en pasos,\nPara seleccionar fecha, hora y cantidad de personas antes de enviar mis datos de contacto.",
            "criterios": "1. El paso 1 debe solicitar fecha, hora y cantidad de personas antes de pedir datos personales.\n2. La fecha mínima seleccionable debe ser la fecha actual y no debe admitirse una fecha pasada.\n3. El horario disponible debe actualizarse dinámicamente según fecha, personas y reservas existentes.\n4. El botón para continuar al paso 2 debe estar deshabilitado cuando no exista mesa disponible.\n5. El paso 2 debe solicitar nombre, teléfono, correo y comentario opcional.\n6. El sistema debe mostrar mensajes claros ante disponibilidad, no disponibilidad o error de conexión.\n7. Al registrar la reserva, debe generarse un código identificador y mostrarse la confirmación al usuario.",
            "reglas": "RN-01: Las reservas públicas solo se aceptan dentro del horario operativo configurado.\nRN-02: La disponibilidad depende de fecha, hora, cantidad de personas y mesas activas.\nRN-03: No se debe solicitar información personal completa si previamente no existe una mesa disponible.",
            "estimacion": "8 puntos de historia / Aprox. 13 horas ideales.",
            "tareas": "T1 (Frontend): Diseñar Reservar.cshtml con bloques de paso 1, paso 2 e indicador de progreso.\nT2 (JavaScript): Implementar navegación, bloqueo de botón y consulta fetch de disponibilidad.\nT3 (Backend): Crear endpoints públicos HorariosDisponibles y Disponibilidad con rate limiting.\nT4 (Validación): Aplicar reglas de fecha vigente, máximo de anticipación y horario operativo.\nT5 (Integración): Asociar cliente existente o crear cliente nuevo durante la reserva.\nT6 (Pruebas): Probar flujo exitoso, horario sin disponibilidad, fecha inválida y validación de campos.",
            "evidencia": "Reserva.Web/Views/Reservas/Reservar.cshtml; Reserva.Web/wwwroot/js/reservas.js; Reserva.Web/Controllers/ReservasController.cs.",
        },
        {
            "id": "HU-06",
            "nombre": "Asignación Automática de Mesa mediante Best-Fit",
            "modulo": "Sprint 3 - Motor de disponibilidad y aforo",
            "prioridad": "Alta",
            "descripcion": "Como sistema de reservas,\nQuiero seleccionar automáticamente la mesa activa de menor capacidad suficiente,\nPara optimizar el aforo del restaurante y evitar desperdicio de mesas grandes.",
            "criterios": "1. El algoritmo debe retornar null si la fecha, hora o cantidad de personas son inválidas.\n2. El sistema debe excluir mesas inactivas antes de evaluar capacidad.\n3. El sistema debe excluir mesas ocupadas por reservas no canceladas en la misma fecha y hora.\n4. La mesa candidata debe tener capacidad mayor o igual a la cantidad de personas solicitada.\n5. Entre varias candidatas, debe seleccionarse la mesa con menor capacidad suficiente.\n6. Si existe empate de capacidad, debe seleccionarse la mesa con menor número.\n7. Las reservas canceladas no deben bloquear la mesa para nuevas solicitudes.",
            "reglas": "RN-01: El aforo se optimiza asignando la mesa más ajustada al tamaño del grupo.\nRN-02: Una reserva cancelada permanece en historial, pero libera la mesa.\nRN-03: El algoritmo no debe permitir doble asignación de una mesa en el mismo horario.",
            "estimacion": "8 puntos de historia / Aprox. 13 horas ideales.",
            "tareas": "T1 (Backend): Implementar FindAvailableMesaAsync en ReservasController.\nT2 (Consulta): Construir subconsulta de mesas ocupadas por fecha, hora y estado.\nT3 (Ordenamiento): Aplicar OrderBy(mesa.Capacidad) y ThenBy(mesa.Numero).\nT4 (Integración): Usar el algoritmo en Disponibilidad, Reservar y validación administrativa.\nT5 (Migración): Ajustar reglas para que reservas canceladas no retengan mesas.\nT6 (Pruebas): Evaluar casos con grupos de 2, 4, 6 y 8 personas, mesas ocupadas y canceladas.",
            "evidencia": "Reserva.Web/Controllers/ReservasController.cs método FindAvailableMesaAsync; Reserva.Infrastructure/Persistence/Migrations/20260626120000_AllowRebookingCancelledTables.cs.",
        },
        {
            "id": "HU-07",
            "nombre": "Consulta, Confirmación y Comprobante de Reserva",
            "modulo": "Sprint 3 - Autoservicio de consulta",
            "prioridad": "Media",
            "descripcion": "Como cliente,\nQuiero consultar el estado de mi reserva y acceder a un comprobante,\nPara confirmar que mi solicitud fue registrada correctamente.",
            "criterios": "1. La pantalla debe permitir elegir consulta por código o por contacto.\n2. Si el método es código, el usuario debe ingresar un código de reserva válido.\n3. Si el método es contacto, el usuario debe ingresar correo o teléfono asociado a reservas existentes.\n4. El resultado debe mostrar código, cliente, fecha, hora, mesa, personas y estado.\n5. El comprobante público debe exigir el código correcto cuando el usuario no está autenticado.\n6. El usuario debe poder cancelar reservas futuras o vigentes desde la consulta, siempre que los datos coincidan.\n7. La pantalla no debe mostrar información de reservas ajenas cuando la validación falla.",
            "reglas": "RN-01: El código de reserva funciona como llave pública de comprobante.\nRN-02: La consulta por contacto puede devolver varias reservas asociadas.\nRN-03: No se permite cancelar desde consulta pública una reserva de fecha pasada.",
            "estimacion": "5 puntos de historia / Aprox. 8 horas ideales.",
            "tareas": "T1 (Backend): Implementar Consultar GET/POST, CancelarConsulta y Comprobante.\nT2 (Seguridad): Validar código de reserva en acceso público a PDF.\nT3 (Servicios): Construir ReservationLookupResult y comprobante mediante ReservationReceiptService.\nT4 (Frontend): Diseñar Consultar.cshtml con resultados múltiples y modal de cancelación.\nT5 (Validación): Normalizar correo, teléfono y código antes de consultar.\nT6 (Pruebas): Verificar búsqueda por código, contacto, descarga de comprobante y cancelación segura.",
            "evidencia": "Reserva.Web/Controllers/ReservasController.cs; Reserva.Web/Views/Reservas/Consultar.cshtml; Reserva.Web/Services/ReservationReceiptService.cs.",
        },
        {
            "id": "HU-08",
            "nombre": "Dashboard Administrativo y Agenda Operativa",
            "modulo": "Sprint 4 - Panel operativo",
            "prioridad": "Alta",
            "descripcion": "Como administrador,\nQuiero visualizar en un panel operativo las reservas pendientes, agenda del día y ocupación,\nPara tomar decisiones rápidas durante la atención del restaurante.",
            "criterios": "1. El dashboard debe mostrar una sección prioritaria de reservas pendientes.\n2. Debe calcular estado operativo en verde, amarillo o rojo según ocupación, mesas inactivas y conflictos.\n3. Debe presentar agenda del día ordenada por hora y número de mesa.\n4. Debe permitir búsqueda global por cliente, código, estado, correo o teléfono.\n5. Debe permitir confirmar o cancelar reservas pendientes desde un modal de acción rápida.\n6. Cada cambio de estado debe persistirse y reflejarse al volver al dashboard.\n7. Los indicadores no deben contabilizar reservas canceladas como ocupación activa.",
            "reglas": "RN-01: Las reservas pendientes tienen prioridad visual sobre el resto de la agenda.\nRN-02: La ocupación actual se calcula con mesas activas y reservas no canceladas.\nRN-03: Las acciones rápidas solo admiten estados Confirmada o Cancelada.",
            "estimacion": "5 puntos de historia / Aprox. 8 horas ideales.",
            "tareas": "T1 (Backend): Extender AdminController con búsqueda global, agenda y métricas.\nT2 (Modelo): Crear OperationalStatusViewModel y ampliar DashboardStatsViewModel.\nT3 (ViewComponent): Actualizar DashboardStatsViewComponent para ocupación y clientes esperados.\nT4 (Frontend): Rediseñar Dashboard.cshtml con panel de estado, prioridades y agenda diaria.\nT5 (Acciones): Implementar CambiarEstadoReserva con validación de estado permitido.\nT6 (Pruebas): Validar conteos, búsqueda, modal de acción, persistencia y redirección.",
            "evidencia": "Reserva.Web/Controllers/AdminController.cs; Reserva.Web/Models/OperationalStatusViewModel.cs; Reserva.Web/Views/Admin/Dashboard.cshtml; Reserva.Web/ViewComponents/DashboardStatsViewComponent.cs.",
        },
        {
            "id": "HU-09",
            "nombre": "Cancelación y Liberación de Mesas",
            "modulo": "Sprint 4 - Gestión de estados de reserva",
            "prioridad": "Alta",
            "descripcion": "Como administrador o cliente autorizado,\nQuiero cancelar una reserva sin eliminar su historial,\nPara liberar la mesa y conservar trazabilidad operativa.",
            "criterios": "1. Una reserva cancelada debe conservarse en la base de datos con su código y datos originales.\n2. Las reservas canceladas deben excluirse del conteo de mesas ocupadas.\n3. Una mesa asociada a reserva cancelada debe estar disponible para otra reserva en la misma fecha y hora.\n4. La cancelación pública debe validar código de reserva y correo asociado.\n5. La cancelación administrativa debe registrar el cambio desde el panel interno.\n6. La agenda y los listados deben mostrar visualmente el estado Cancelada.\n7. El sistema debe evitar cancelaciones inconsistentes o sobre reservas inexistentes.",
            "reglas": "RN-01: Cancelar no equivale a eliminar; la reserva permanece como historial operativo.\nRN-02: La disponibilidad solo considera como ocupadas las reservas pendientes o confirmadas.\nRN-03: El personal administrativo puede cancelar desde el dashboard y el cliente desde consulta validada.",
            "estimacion": "5 puntos de historia / Aprox. 8 horas ideales.",
            "tareas": "T1 (Migración): Crear AllowRebookingCancelledTables para documentar la regla de liberación.\nT2 (Backend): Excluir Cancelada en validaciones de conflictos y disponibilidad.\nT3 (Consulta pública): Implementar CancelarConsulta con validación de código y correo.\nT4 (Dashboard): Integrar cancelación desde CambiarEstadoReserva.\nT5 (Frontend): Aplicar badges y estilos para diferenciar estados.\nT6 (Pruebas): Confirmar que una mesa cancelada puede reservarse nuevamente.",
            "evidencia": "Reserva.Infrastructure/Persistence/Migrations/20260626120000_AllowRebookingCancelledTables.cs; Reserva.Web/Controllers/ReservasController.cs; Reserva.Web/Controllers/AdminController.cs.",
        },
        {
            "id": "HU-10",
            "nombre": "Notificación y Redirección a WhatsApp",
            "modulo": "Sprint 4 - Comunicación con el cliente",
            "prioridad": "Media",
            "descripcion": "Como personal operativo,\nQuiero generar un enlace de WhatsApp con los datos de la reserva,\nPara comunicar al cliente la confirmación, cambio o seguimiento de su solicitud.",
            "criterios": "1. El sistema debe construir una URL de WhatsApp con número de cliente y mensaje prellenado.\n2. El mensaje debe incluir código de reserva, nombre del cliente, fecha, hora, mesa y estado.\n3. La URL debe codificar correctamente espacios, tildes y caracteres especiales.\n4. El enlace debe generarse desde el servicio de notificación y no directamente desde la vista.\n5. El enlace debe estar disponible en confirmación o consulta cuando exista teléfono válido.\n6. La notificación no debe reemplazar la persistencia interna ni el comprobante de reserva.\n7. Si el teléfono no es válido, el sistema debe evitar generar un enlace incorrecto.",
            "reglas": "RN-01: WhatsApp es un canal auxiliar de comunicación, no la fuente oficial de estado.\nRN-02: La reserva oficial se mantiene en la base de datos con su código y estado.\nRN-03: El mensaje enviado debe ser coherente con el estado vigente de la reserva.",
            "estimacion": "3 puntos de historia / Aprox. 5 horas ideales.",
            "tareas": "T1 (Servicio): Implementar BuildWhatsAppUrl y normalización de teléfono en ReservationNotificationService.\nT2 (Modelo): Exponer WhatsAppUrl en ReservationConfirmationViewModel y resultados de consulta.\nT3 (Frontend): Agregar botón de contacto en vistas de confirmación y consulta.\nT4 (Codificación): Validar Encode de mensaje para URL segura.\nT5 (Integración): Conectar actualización de estado con mensajes de seguimiento.\nT6 (Pruebas): Probar enlace con teléfono válido, caracteres especiales y estados distintos.",
            "evidencia": "Reserva.Web/Services/ReservationNotificationService.cs; Reserva.Web/Models/ReservationConfirmationViewModel.cs; Reserva.Web/Views/Reservas/Confirmacion.cshtml; Reserva.Web/Views/Reservas/Consultar.cshtml.",
        },
    ]

    for index, story in enumerate(detailed_stories, start=1):
        add_user_story_detail(doc, annex_table_counter, index, story)

    add_page_break()
    add_heading(doc, "Anexo 5. Entregables de Ingeniería de Software", 3)
    add_paragraph(
        doc,
        "Los entregables técnicos permiten evidenciar la construcción progresiva de la plataforma, desde el diseño arquitectónico hasta la validación funcional del MVP."
    )
    deliverable_rows = [
        ("Arquitectura N-capas", "Solución ASP.NET Core MVC organizada en capas Domain, Infrastructure, Web y Tests.", "Reserva.slnx y proyectos Reserva.Domain, Reserva.Infrastructure, Reserva.Web, Reserva.Tests."),
        ("Modelo de datos", "Entidades Cliente, Reserva, Mesa, Plato y Usuario, persistidas mediante Entity Framework Core.", "Reserva.Infrastructure/Persistence/ReservationDbContext.cs."),
        ("Seguridad", "Hashing PBKDF2 para almacenamiento seguro de contraseñas administrativas.", "Reserva.Infrastructure/Security/PasswordHashing.cs."),
        ("Aforo Best-Fit", "Asignación automática de la mesa de menor capacidad suficiente para una reserva.", "Reserva.Web/Controllers/ReservasController.cs."),
        ("Wizard de reserva", "Formulario público dividido en pasos para disponibilidad y datos de contacto.", "Reserva.Web/Views/Reservas/Reservar.cshtml y wwwroot/js/reservas.js."),
        ("Panel operativo", "Dashboard con métricas, búsqueda global, agenda diaria y acciones rápidas.", "Reserva.Web/Controllers/AdminController.cs y Views/Admin/Dashboard.cshtml."),
    ]
    add_table(doc, ["Entregable", "Descripción", "Evidencia técnica"], deliverable_rows, [4.0, 6.2, 6.6])

    add_heading(doc, "Anexo 6. Evidencia de Control de Versiones", 3)
    git_rows = [
        ("Initial commit: Configuración de la solución Mikuy", "Configuración inicial de la solución y estructura base del repositorio."),
        ("feat: Estructura base en capas (.NET Core, Entity Framework)", "Formalización de la arquitectura por capas y persistencia."),
        ("Merge branch 'feature/best-fit' into main (Aforo Best-Fit)", "Integración del algoritmo de asignación de mesas por capacidad óptima."),
        ("feat: Implementación de wizard de reservas y redirección a WhatsApp", "Implementación del flujo público de reserva y comunicación."),
        ("Merge branch 'feature/panel-operativo' into main (Dashboard and CRUDs)", "Integración del panel administrativo, agenda y CRUDs principales."),
    ]
    add_table(doc, ["Commit / integración", "Descripción"], git_rows, [7.0, 9.8])
    add_note(doc, "La trazabilidad del repositorio respalda la evolución incremental del producto durante los sprints. Elaboración propia.")

    doc.save(OUT_DOC)


if __name__ == "__main__":
    build()
    print(OUT_DOC)
